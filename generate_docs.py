"""Generate PyVizBuilder Production Deployment Guide as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


# ── Helpers ───────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x25, 0x30, 0x47)
        run.font.size = Pt(14)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        run.font.size = Pt(12)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=10):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code_block(doc, code: str, caption: str = ""):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.runs[0].font.size  = Pt(8)
        cp.runs[0].font.italic = True
        cp.runs[0].font.color.rgb = RGBColor(0x66, 0x70, 0x85)
        cp.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1E1E2E")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run  = para.add_run(code)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xC9, 0xD1, 0xD9)
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_info_box(doc, text: str, bg="EFF6FF", fg=(0x1D, 0x4E, 0xD8)):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell  = table.cell(0, 0)
    set_cell_bg(cell, bg)
    para  = cell.paragraphs[0]
    run   = para.add_run(text)
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(*fg)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_warning_box(doc, text: str):
    add_info_box(doc, f"⚠  {text}", bg="FFFBEB", fg=(0x92, 0x40, 0x0E))


def add_step_table(doc, rows_data: list):
    """rows_data = list of (step_no, title, detail) tuples."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, ["#", "Step", "Detail"]):
        set_cell_bg(cell, "253047")
        run = cell.paragraphs[0].add_run(txt)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9)

    for no, title, detail in rows_data:
        row  = tbl.add_row().cells
        row[0].text = str(no)
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = title
        row[1].paragraphs[0].runs[0].font.bold = True
        row[1].paragraphs[0].runs[0].font.size = Pt(9)
        row[2].text = detail
        row[2].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(row[0], "F8FAFC")

    tbl.columns[0].width = Cm(1)
    tbl.columns[1].width = Cm(5)
    tbl.columns[2].width = Cm(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_kv_table(doc, rows_data: list, headers=("Setting", "Value / Description")):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, headers):
        set_cell_bg(cell, "253047")
        run = cell.paragraphs[0].add_run(txt)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9)

    for i, (key, val) in enumerate(rows_data):
        row  = tbl.add_row().cells
        bg   = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(row[0], bg); set_cell_bg(row[1], bg)
        kr = row[0].paragraphs[0].add_run(key)
        kr.font.name = "Courier New"
        kr.font.size = Pt(9)
        vr = row[1].paragraphs[0].add_run(val)
        vr.font.size = Pt(9)

    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ═════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width   = Inches(8.5)
section.page_height  = Inches(11)
section.left_margin  = Inches(1)
section.right_margin = Inches(1)
section.top_margin   = Inches(1)
section.bottom_margin = Inches(1)

# ── Cover ──────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_r = title_p.add_run("PyVizBuilder")
title_r.font.size  = Pt(28)
title_r.font.bold  = True
title_r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_r = sub_p.add_run("Production Deployment Guide")
sub_r.font.size  = Pt(16)
sub_r.font.color.rgb = RGBColor(0x66, 0x70, 0x85)

doc.add_paragraph()

add_info_box(doc,
    "Scope: this guide covers deploying PyVizBuilder against live BigQuery tables.\n"
    "Mock mode (USE_MOCK = True) and the chart gallery (--test) are explicitly excluded.",
    bg="F0F9FF", fg=(0x03, 0x69, 0xA1))

doc.add_page_break()

# ── 1. Prerequisites ──────────────────────────────────────────────────
add_heading(doc, "1. Prerequisites")

add_heading(doc, "1.1  Python & packages", level=2)
add_para(doc, "Python 3.10 or higher required. Install all dependencies:")
add_code_block(doc, "pip install pandas numpy matplotlib pillow altair vl-convert-python google-cloud-bigquery google-auth",
               "Install dependencies")

add_heading(doc, "1.2  GCP project requirements", level=2)
add_kv_table(doc, [
    ("GCP Project",          "An active Google Cloud project with billing enabled."),
    ("BigQuery API",         "Enable at: GCP Console → APIs & Services → BigQuery API."),
    ("Service Account",      "A service account with a downloaded JSON key file."),
    ("IAM roles (minimum)",  "roles/bigquery.dataViewer\nroles/bigquery.dataEditor (for email_output writes)\nroles/bigquery.jobUser"),
])

add_heading(doc, "1.3  BigQuery datasets & tables", level=2)
add_para(doc, "Three BigQuery objects must exist before the engine runs:")
add_kv_table(doc, [
    ("email_list",          "Table — one row per (report_name, recipient) dispatch."),
    ("chart_config_view",   "VIEW (UNION ALL of SELECT rows) — one row per chart config."),
    ("email_output",        "Table — written by the engine after each run."),
    ("prod_reports.*",      "Source data tables referenced in chart bq_table fields."),
])

# ── 2. GCP Service Account Setup ─────────────────────────────────────
add_heading(doc, "2. GCP Service Account Setup")

add_step_table(doc, [
    (1, "Open GCP Console", "https://console.cloud.google.com"),
    (2, "Navigate to IAM", "IAM & Admin → Service Accounts"),
    (3, "Create account", "Click 'Create Service Account'. Name it e.g. pyvizbuilder-prod."),
    (4, "Assign roles", "Grant the three roles listed in section 1.2."),
    (5, "Create JSON key", "Service account → Keys tab → Add Key → Create new key → JSON."),
    (6, "Save key file", "Save downloaded JSON as  service_account_key.json  in the same folder as chart_email_engine_v15.py."),
])

add_warning_box(doc,
    "Never commit service_account_key.json to version control. Add it to .gitignore immediately.")

add_code_block(doc, "echo service_account_key.json >> .gitignore", ".gitignore entry")

# ── 3. BigQuery Schema Setup ──────────────────────────────────────────
add_heading(doc, "3. BigQuery Schema Setup")

add_heading(doc, "3.1  Run DDL + seed data", level=2)
add_para(doc, "Two SQL files ship with the engine. Run them in order:")
add_code_block(doc,
    "# Step 1 — create tables and view (DDL only)\n"
    "bq query --use_legacy_sql=false < bigquery_setup_v15_schema.sql\n\n"
    "# Step 2 — insert seed rows and run verification queries\n"
    "bq query --use_legacy_sql=false < bigquery_setup_v15_data.sql",
    "Run from the PyViz/ directory")

add_heading(doc, "3.2  email_list table schema", level=2)
add_kv_table(doc, [
    ("email_id",       "STRING — unique identifier for this dispatch (used as filename suffix)."),
    ("report_name",    "STRING — joins to chart_config_view.report_name."),
    ("recipient_email","STRING — destination address (informational; engine does not send email)."),
    ("subject",        "STRING — email subject line. Supports {token} substitution."),
    ("html_template",  "STRING — complete HTML document. Must contain {{VARIABLE_NAME}} placeholders."),
    ("filter_params",  "JSON STRING or NULL — per-recipient variable overrides (deprecated in v15; use email_id)."),
], headers=("Column", "Description"))

add_heading(doc, "3.3  chart_config_view schema", level=2)
add_para(doc, "This is a VIEW — never INSERT into it. Add charts by appending UNION ALL blocks to the view DDL.")
add_kv_table(doc, [
    ("report_name",    "STRING — must match email_list.report_name exactly."),
    ("sort_position",  "INTEGER — render order within the report."),
    ("variable_name",  "STRING — must match {{VARIABLE_NAME}} placeholder in html_template (uppercase)."),
    ("chart_type",     "STRING — one of 8 Altair types (see section 5)."),
    ("bq_table",       "STRING — full path: project.dataset.table. Supports {token} substitution."),
    ("filters",        "STRING — SQL WHERE clause fragment. Supports {token} substitution."),
    ("x_column",       "STRING — column name for x-axis / label."),
    ("y_columns",      "STRING — comma-separated column names for value(s)."),
    ("legend",         "STRING — 'yes' or 'no'."),
    ("title",          "STRING — chart title. Supports {token} substitution."),
    ("subtitle",       "STRING — chart subtitle."),
    ("color_theme",    "STRING — blue | teal | warm | cool | vibrant | purple | slate | green | default | rainbow."),
    ("show_values",    "STRING — 'yes' or 'no'."),
    ("sort_order",     "STRING — desc | asc | none."),
    ("width_px",       "INTEGER — chart width in pixels."),
    ("height_px",      "INTEGER — chart height in pixels."),
    ("ref_line_value", "STRING — optional reference line value. Supports {token}."),
    ("ref_line_label", "STRING — label for reference line."),
    ("x_label",        "STRING — x-axis label."),
    ("y_label",        "STRING — y-axis label."),
    ("dark_mode",      "STRING — 'yes' or 'no'."),
    ("hue_column",     "STRING — column for color grouping / long-form series."),
], headers=("Column", "Description"))

add_heading(doc, "3.4  email_output table schema", level=2)
add_para(doc, "Written automatically by the engine. Do not create rows manually.")
add_kv_table(doc, [
    ("email_id",        "STRING"),
    ("report_name",     "STRING"),
    ("recipient_email", "STRING"),
    ("subject",         "STRING"),
    ("final_html",      "STRING — complete rendered HTML with base64 chart images injected."),
    ("charts_injected", "INTEGER"),
    ("total_charts",    "INTEGER"),
    ("status",          "STRING — SUCCESS | WARN | FAILED"),
    ("error_message",   "STRING — empty on SUCCESS"),
    ("processed_at",    "TIMESTAMP"),
], headers=("Column", "Type / Notes"))

# ── 4. Configuration ──────────────────────────────────────────────────
add_heading(doc, "4. Configuration")

add_para(doc, "All settings live in PyVizBuilder/config.py. Edit this file before running in production.")

add_heading(doc, "4.1  Required settings", level=2)
add_code_block(doc,
    "# PyVizBuilder/config.py\n\n"
    "SERVICE_ACCOUNT_KEY_FILE = \"service_account_key.json\"  # path to JSON key\n"
    "PROJECT_ID               = \"my-gcp-project\"            # GCP project ID\n\n"
    "EMAIL_LIST_TABLE         = \"my_project.reporting.email_list\"\n"
    "CHART_CONFIG_VIEW        = \"my_project.reporting.chart_config_view\"\n"
    "EMAIL_OUTPUT_TABLE       = \"my_project.reporting.email_output\"\n\n"
    "USE_MOCK                 = False   # MUST be False for production\n"
    "WRITE_MODE               = \"APPEND\"   # APPEND keeps history; TRUNCATE overwrites")

add_heading(doc, "4.2  Optional settings", level=2)
add_kv_table(doc, [
    ("MAX_WORKERS = 4",       "Parallel chart-render threads per email. Increase on multi-core servers."),
    ("OUTPUT_DIR = 'output_emails'", "Local directory for rendered HTML files."),
    ("TABLE_VARS = {'env': 'prod'}", "Global {token} overrides applied to every email/chart in the run."),
    ("WRITE_MODE = 'TRUNCATE'",     "Use TRUNCATE to replace email_output on each run instead of appending."),
])

add_heading(doc, "4.3  TABLE_VARS — global token overrides", level=2)
add_para(doc,
    "TABLE_VARS injects tokens into every email and chart SQL filter in the run. "
    "Use for environment, region, or fiscal-year values that apply globally.")
add_code_block(doc,
    "TABLE_VARS: dict = {\n"
    "    \"env\":    \"prod\",\n"
    "    \"region\": \"APAC\",\n"
    "    \"fy\":     \"2026\",\n"
    "}")

add_info_box(doc,
    "Priority order (lowest → highest):\n"
    "  1. Built-in date tokens  ({today}, {this_month}, {this_quarter_year}, etc.)\n"
    "  2. TABLE_VARS\n"
    "  3. email_id  (from email_list — per-dispatch identifier)")

# ── 5. Supported Chart Types ──────────────────────────────────────────
add_heading(doc, "5. Supported Chart Types")

add_kv_table(doc, [
    ("line_altair",    "Multi-series line with point overlay. hue_column for long-form grouping."),
    ("bar_altair",     "Sorted horizontal bar. sort_order: desc | asc | none."),
    ("scatter_altair", "Bubble scatter. y_columns[1] = size column (optional). hue_column for color."),
    ("heatmap_altair", "Rect heatmap with sequential color scale. hue_column = row axis."),
    ("area_altair",    "Stacked area. Multiple y_columns melted to long-form automatically."),
    ("strip_altair",   "Strip / jitter plot with median tick overlay."),
    ("boxplot_altair", "Box-and-whisker with outlier marks. hue_column for grouped boxes."),
    ("arc_altair",     "Pie or donut arc. Set y_columns[1] = 'donut' for donut style."),
], headers=("chart_type", "Description / Key columns"))

# ── 6. Built-in Date Tokens ───────────────────────────────────────────
add_heading(doc, "6. Built-in Date Tokens")
add_para(doc, "Available in bq_table paths, filters, titles, subtitles, and email subjects.")

add_kv_table(doc, [
    ("{today}",              "Current date: 2026-05-01"),
    ("{yesterday}",          "Previous date"),
    ("{tomorrow}",           "Next date"),
    ("{this_month}",         "YYYY-MM format: 2026-05"),
    ("{this_month_name}",    "Full month name: May"),
    ("{last_month}",         "YYYY-MM of previous month"),
    ("{this_quarter}",       "Q2"),
    ("{this_quarter_year}",  "Q2-2026"),
    ("{last_quarter_year}",  "Q1-2026"),
    ("{this_half}",          "H1 or H2"),
    ("{this_year}",          "2026"),
    ("{last_year}",          "2025"),
    ("{this_week}",          "ISO week: W18"),
    ("{env}",                "'prod' by default; override via TABLE_VARS"),
], headers=("Token", "Example / Notes"))

add_para(doc, "Token in SQL filter example:")
add_code_block(doc,
    "-- In chart_config_view filters column:\n"
    "report_month = '{this_month}' AND region = '{region}'\n\n"
    "-- Resolves to:\n"
    "report_month = '2026-05' AND region = 'APAC'")

# ── 7. Running in Production ──────────────────────────────────────────
add_heading(doc, "7. Running in Production")

add_heading(doc, "7.1  Single run", level=2)
add_code_block(doc,
    "# From the PyViz/ directory\n"
    "python chart_email_engine_v15.py",
    "Production run command")

add_para(doc, "The engine will:")
for item in [
    "Authenticate with GCP using service_account_key.json.",
    "Build runtime variable context (date tokens + TABLE_VARS).",
    "JOIN email_list ⋈ chart_config_view on report_name.",
    "Group rows by (report_name, email_id).",
    "Render all charts in parallel (MAX_WORKERS threads).",
    "Inject base64 PNG images into html_template placeholders.",
    "Write one HTML file per (report_name, email_id) to output_emails/.",
    "Write one row per (report_name, email_id) to email_output table.",
]:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(9)

add_heading(doc, "7.2  Scheduled / cron run", level=2)
add_code_block(doc,
    "# Linux/Mac crontab — run daily at 06:00\n"
    "0 6 * * * cd /opt/pyviz && python chart_email_engine_v15.py >> logs/run.log 2>&1\n\n"
    "# Windows Task Scheduler action\n"
    "Program: python\n"
    "Arguments: chart_email_engine_v15.py\n"
    "Start in: C:\\deploy\\PyViz")

add_heading(doc, "7.3  Output files", level=2)
add_kv_table(doc, [
    ("output_emails/<report>__<email_id>.html",
     "Self-contained HTML with all charts as inline base64 PNGs. Safe to attach or serve directly."),
    ("email_output BQ table",
     "Row per dispatch: final_html, charts_injected, total_charts, status, error_message, processed_at."),
])

# ── 8. Variable Resolution in Depth ──────────────────────────────────
add_heading(doc, "8. Variable Resolution Reference")

add_para(doc, "Token syntax: single braces {token} in SQL, titles, filters. "
              "Double braces {{VARIABLE_NAME}} in html_template for chart injection.")

add_info_box(doc,
    "Resolution order (lowest → highest priority):\n"
    "  1. Built-in tokens (date/time)\n"
    "  2. TABLE_VARS (global, set in config.py)\n"
    "  3. email_id (per-dispatch, from email_list)\n\n"
    "Higher-priority values override lower-priority ones with the same key name.")

add_heading(doc, "8.1  Per-dispatch email_id example", level=2)
add_para(doc,
    "If email_list has email_id = 'region_APAC', the engine merges "
    "{email_id: 'region_APAC'} into the variable context. "
    "Filters can then use {email_id} to scope SQL results per recipient.")

add_code_block(doc,
    "-- chart_config_view.filters:\n"
    "region = '{email_id}'\n\n"
    "-- email_list row with email_id = 'APAC':\n"
    "-- resolves to:  region = 'APAC'\n\n"
    "-- email_list row with email_id = 'EMEA':\n"
    "-- resolves to:  region = 'EMEA'")

# ── 9. chart_config_view Maintenance ─────────────────────────────────
add_heading(doc, "9. Adding New Charts (chart_config_view)")

add_para(doc,
    "chart_config_view is a BigQuery VIEW defined as UNION ALL of SELECT rows. "
    "Never INSERT into it. To add a chart, alter the view DDL.")

add_code_block(doc,
    "-- Append to the existing view definition:\n"
    "CREATE OR REPLACE VIEW `my_project.reporting.chart_config_view` AS\n\n"
    "  -- existing rows ...\n\n"
    "  UNION ALL\n\n"
    "  SELECT\n"
    "    'monthly_report'   AS report_name,\n"
    "    99                 AS sort_position,\n"
    "    'NEW_CHART'        AS variable_name,\n"
    "    'bar_altair'       AS chart_type,\n"
    "    'proj.ds.my_table' AS bq_table,\n"
    "    'month = \\'{this_month}\\'' AS filters,\n"
    "    'category'         AS x_column,\n"
    "    'revenue'          AS y_columns,\n"
    "    'yes'              AS legend,\n"
    "    'Revenue by Category' AS title,\n"
    "    '{this_month_name} {this_year}' AS subtitle,\n"
    "    'teal'             AS color_theme,\n"
    "    'yes'              AS show_values,\n"
    "    'desc'             AS sort_order,\n"
    "    620                AS width_px,\n"
    "    320                AS height_px,\n"
    "    ''                 AS ref_line_value,\n"
    "    ''                 AS ref_line_label,\n"
    "    ''                 AS x_label,\n"
    "    'Revenue ($)'      AS y_label,\n"
    "    'no'               AS dark_mode,\n"
    "    ''                 AS hue_column",
    "Add chart to view DDL")

add_para(doc,
    "Then add {{NEW_CHART}} placeholder in the html_template column of the corresponding email_list row.")

# ── 10. html_template Structure ───────────────────────────────────────
add_heading(doc, "10. html_template Structure")

add_para(doc,
    "html_template in email_list owns the complete HTML document. "
    "The engine only replaces {{PLACEHOLDER}} tokens with <img> tags. "
    "It does not add DOCTYPE, head, styles, or body tags.")

add_code_block(doc,
    "<!DOCTYPE html>\n"
    "<html lang=\"en\">\n"
    "<head>\n"
    "  <meta charset=\"UTF-8\" />\n"
    "  <title>{subject}</title>\n"
    "  <style>\n"
    "    body { font-family: Arial, sans-serif; background: #f4f6fa; }\n"
    "    .chart-wrap { margin: 24px auto; max-width: 680px; }\n"
    "  </style>\n"
    "</head>\n"
    "<body>\n"
    "  <h1>Monthly Report — {this_month_name} {this_year}</h1>\n"
    "  <div class=\"chart-wrap\">{{REVENUE_TREND}}</div>\n"
    "  <div class=\"chart-wrap\">{{REGION_BAR}}</div>\n"
    "  <div class=\"chart-wrap\">{{MARKET_SHARE}}</div>\n"
    "</body>\n"
    "</html>",
    "Example html_template value (stored in email_list.html_template)")

add_warning_box(doc,
    "{{PLACEHOLDER}} names must be UPPERCASE and match variable_name in chart_config_view exactly. "
    "Mismatch = placeholder left as-is in output HTML (no error thrown).")

# ── 11. Monitoring & Troubleshooting ─────────────────────────────────
add_heading(doc, "11. Monitoring & Troubleshooting")

add_heading(doc, "11.1  Status codes in email_output", level=2)
add_kv_table(doc, [
    ("SUCCESS", "All {{PLACEHOLDER}} tokens in the template were rendered and injected."),
    ("WARN",    "Partial render: some placeholders had no matching config row or returned 0 rows."),
    ("FAILED",  "No charts injected, or a fatal error occurred. Check error_message column."),
], headers=("Status", "Meaning"))

add_heading(doc, "11.2  Common errors", level=2)
add_kv_table(doc, [
    ("FileNotFoundError: service_account_key.json",
     "Key file not found. Save the GCP JSON key as service_account_key.json in the PyViz/ directory."),
    ("403 Permission denied on BigQuery",
     "Service account missing required IAM roles. See section 1.2."),
    ("0 rows returned for <chart>",
     "SQL returned empty result. Check bq_table path, filters, and that source data exists for the token-resolved date range."),
    ("'<VAR>' not in chart_config_view",
     "Placeholder {{VAR}} found in html_template but no chart_config_view row has variable_name = 'VAR' for that report_name."),
    ("Variable resolution failed",
     "A {token} in bq_table or filters could not be resolved. Verify TABLE_VARS or email_id supplies the token."),
    ("Unknown chart_type='...'",
     "chart_type value not in the 8 supported types. Check spelling."),
], headers=("Error", "Fix"))

add_heading(doc, "11.3  Debug SQL queries", level=2)
add_para(doc, "To print resolved SQL before execution, add a print() call in PyVizBuilder/engine.py:")
add_code_block(doc,
    "# In _render_one(), after sql = build_select(cfg):\n"
    "print(f'  [DEBUG SQL] {var}:\\n{sql}\\n')")

add_heading(doc, "11.4  Verify BQ write", level=2)
add_code_block(doc,
    "bq query --use_legacy_sql=false \\\n"
    "  'SELECT report_name, email_id, status, charts_injected, processed_at\n"
    "   FROM `my_project.reporting.email_output`\n"
    "   ORDER BY processed_at DESC\n"
    "   LIMIT 20'")

# ── 12. File Structure Reference ─────────────────────────────────────
add_heading(doc, "12. File & Module Reference")

add_code_block(doc,
    "PyViz/\n"
    "├── chart_email_engine_v15.py   ← entry point (run this)\n"
    "├── service_account_key.json    ← GCP key (never commit)\n"
    "├── bigquery_setup_v15_schema.sql\n"
    "├── bigquery_setup_v15_data.sql\n"
    "├── output_emails/              ← rendered HTML output\n"
    "└── PyVizBuilder/\n"
    "    ├── config.py               ← ALL settings edited here\n"
    "    ├── vars.py                 ← token resolution engine\n"
    "    ├── bq_client.py            ← BigQuery auth + query/write\n"
    "    ├── sql_builder.py          ← SELECT builder + config parser\n"
    "    ├── engine.py               ← main pipeline (process_emails)\n"
    "    ├── mock_data.py            ← synthetic data (dev only)\n"
    "    ├── test_gallery.py         ← HTML gallery (dev only)\n"
    "    └── charts/\n"
    "        ├── altair_helpers.py   ← colour palettes, legend composer\n"
    "        ├── altair_renderers.py ← 8 chart render functions\n"
    "        └── __init__.py         ← render_chart() dispatcher")

# ── 13. Production Checklist ──────────────────────────────────────────
add_heading(doc, "13. Pre-Launch Checklist")

add_step_table(doc, [
    (1,  "USE_MOCK = False",            "Set in PyVizBuilder/config.py"),
    (2,  "PROJECT_ID",                  "Set to your GCP project ID"),
    (3,  "Table paths",                 "EMAIL_LIST_TABLE, CHART_CONFIG_VIEW, EMAIL_OUTPUT_TABLE all point to correct tables"),
    (4,  "service_account_key.json",    "File present in PyViz/ directory"),
    (5,  "IAM roles",                   "Service account has dataViewer, dataEditor, jobUser"),
    (6,  "Schema deployed",             "bigquery_setup_v15_schema.sql executed successfully"),
    (7,  "email_list populated",        "At least one row with valid html_template and {{PLACEHOLDER}} tokens"),
    (8,  "chart_config_view populated", "VIEW has rows for every report_name used in email_list"),
    (9,  "Source tables exist",         "All bq_table values in chart_config_view resolve to existing tables"),
    (10, ".gitignore",                  "service_account_key.json listed"),
    (11, "Test run",                    "python chart_email_engine_v15.py — verify status = SUCCESS in email_output"),
])

# ── Save ──────────────────────────────────────────────────────────────
out = "PyVizBuilder_Production_Deployment_Guide.docx"
doc.save(out)
print(f"Saved: {out}")
